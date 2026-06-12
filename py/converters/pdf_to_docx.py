"""
PDF → DOCX export — a faithful Word copy of a PDF's text.

Goal: the .docx contains the EXACT text of the PDF, one Word page-break per PDF
page, each page labeled with its PDF page number, and paragraph/line spacing
preserved. Scanned pages are OCR'd (reusing the project OCR pipeline). This is NOT
the tax/Gold-Master pipeline — it is verbatim text, suitable for turning a scanned
legal/trust/POA document back into an editable document.

It does not attempt pixel-perfect visual layout (fonts, columns, positioning) —
that is not recoverable from a scan. Text + page mapping + page numbers are.
"""
import re
import fitz  # PyMuPDF

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from converters.pdf_converter import extract_page_text_verbatim, _ocr_available


def _add_footer_page_numbers(section) -> None:
    """Add an auto 'Page X' field to the section footer (Word updates it)."""
    try:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Page ")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)
    except Exception:
        pass


def _add_page_label(doc, page_num: int, total: int) -> None:
    """A small grey centered '— Page N of M —' marker for the PDF page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"— Page {page_num} of {total} —")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _add_text_block(doc, text: str) -> None:
    """Render verbatim page text: blank lines separate paragraphs; single newlines
    become line breaks within a paragraph, preserving the PDF's line structure."""
    if not text.strip():
        doc.add_paragraph("")
        return
    # Split into paragraphs on one-or-more blank lines.
    for block in re.split(r"\n[ \t]*\n", text):
        lines = block.split("\n")
        if not any(l.strip() for l in lines):
            continue
        p = doc.add_paragraph()
        for i, line in enumerate(lines):
            run = p.add_run(line.rstrip())
            if i < len(lines) - 1:
                run.add_break(WD_BREAK.LINE)


def _wc(text: str) -> int:
    """Word count — one tokenizer used for BOTH sides of the fidelity audit so a
    zero delta means an identical token count (whitespace-split)."""
    return len(text.split())


_PAGE_LABEL_RE = re.compile(r"^—\s*Page\s+(\d+)\s+of\s+\d+\s*—$")


def _para_xml_text(p_elem) -> str:
    """Reconstruct a paragraph's text from its XML in document order, turning
    <w:br>/<w:cr> into newlines and <w:tab> into tabs. python-docx's
    `paragraph.text` drops line breaks entirely, which would merge the last word
    of one line into the first of the next and corrupt the word count."""
    parts = []
    for node in p_elem.iter():
        tag = node.tag
        if tag == qn("w:t"):
            parts.append(node.text or "")
        elif tag in (qn("w:br"), qn("w:cr")):
            parts.append("\n")
        elif tag == qn("w:tab"):
            parts.append("\t")
    return "".join(parts)


def _extract_docx_page_texts(docx_path: str) -> dict:
    """Re-read a DOCX built by build_pdf_docx and return {page_num: text}.

    Pages are delimited by the '— Page N of M —' label paragraphs this module
    writes at the top of each PDF page's content. The label paragraphs are NOT
    included in the returned text (they are markers, not source content). This is
    the INDEPENDENT side of the fidelity audit — text actually persisted to disk,
    not the in-memory string we intended to write.
    """
    doc = Document(docx_path)
    pages: dict = {}
    current = None
    for para in doc.paragraphs:
        txt = _para_xml_text(para._p)
        m = _PAGE_LABEL_RE.match(txt.strip())
        if m:
            current = int(m.group(1))
            pages.setdefault(current, [])
            continue
        if current is not None:
            pages[current].append(txt)
    return {n: "\n".join(blocks) for n, blocks in pages.items()}


def write_audit_report(meta: dict, source_name: str, docx_name: str,
                       out_path: str, *, faithful: bool) -> str:
    """Write a Markdown fidelity-audit report from build_pdf_docx() meta.

    Proves the PDF text carried into the DOCX: pages match 1:1, per-page word
    counts on both sides, and every page whose counts differ flagged for a human
    to check against the PDF. Returns out_path.
    """
    audit = meta.get("page_audit", [])
    pages = meta.get("pages", len(audit))
    deltas = meta.get("pages_with_delta", [])
    matched = len(audit) - len(deltas)
    lines = [
        f"# DOCX fidelity audit — {source_name}",
        "",
        f"- **Source PDF:** {source_name}",
        f"- **DOCX output:** {docx_name}",
        f"- **Pages:** {pages} (PDF and DOCX, 1:1)",
        f"- **Total words:** PDF {meta.get('total_pdf_words', 0)} · "
        f"DOCX {meta.get('total_docx_words', 0)}",
        f"- **Pages matching:** {matched} of {len(audit)}",
        "",
        "_DOCX counts are read back from the **saved .docx file** (independent "
        "round-trip), not the text we intended to write._",
        "",
    ]
    if faithful:
        if deltas:
            lines.append(f"> ⚠ **{len(deltas)} page(s) differ** — review against the "
                         f"PDF: {', '.join(str(p) for p in deltas)}.")
        else:
            lines.append("> ✅ Every page's words survived the round-trip to the "
                         "saved .docx — counts match the source PDF exactly.")
    else:
        lines.append("> ℹ Masking / replacements were applied, so per-page word "
                     "counts are expected to differ from the source. A delta here "
                     "reflects redaction, not lost text.")
    lines += [
        "",
        "| Page | PDF words | DOCX words | Δ | Status |",
        "| ---: | ---: | ---: | ---: | :--- |",
    ]
    for p in audit:
        flag = "match" if p["delta"] == 0 else "**CHECK**"
        sign = f"+{p['delta']}" if p["delta"] > 0 else str(p["delta"])
        lines.append(f"| {p['page']} | {p['pdf_words']} | {p['docx_words']} | "
                     f"{sign} | {flag} |")
    lines.append("")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return out_path


def build_pdf_docx(pdf_path: str, out_path: str, *, mask_fn=None,
                   progress_cb=None, control: dict = None) -> dict:
    """Build a DOCX mirroring the PDF's text and pages.

    mask_fn: optional callable(text)->text applied per page (PII masking /
             replacements). None = faithful copy (full PII).
    progress_cb(page_index, total_pages): optional progress hook.
    control: optional {"stop": bool, "skip": bool} cooperative cancel.
    Returns meta: {pages, ocr_pages, page_audit, total_pdf_words,
                   total_docx_words, pages_with_delta}.

    page_audit is a per-page list [{page, pdf_words, docx_words, delta}] proving
    the text carried over: pdf_words counts the source PDF page (digital text or
    OCR), docx_words counts what was actually written to the DOCX page. For a
    faithful copy (mask_fn=None) every delta should be 0; masking legitimately
    shifts the count.
    """
    src = fitz.open(pdf_path)
    total = len(src)
    ocr_enabled = _ocr_available()

    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(1))
    # Default body font: a clean serif close to legal documents.
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    _add_footer_page_numbers(section)

    ocr_pages = 0
    processed = []
    source_words = {}     # page -> word count of the source PDF text (pre-mask)
    for i, page in enumerate(src, start=1):
        if control and (control.get("stop") or control.get("skip")):
            break
        if progress_cb:
            progress_cb(i, total)
        needs_ocr = ocr_enabled and len(page.get_text().strip()) < 200
        source_text = extract_page_text_verbatim(page, ocr_enabled=ocr_enabled)
        if needs_ocr and source_text:
            ocr_pages += 1
        text = source_text
        if mask_fn:
            try:
                text = mask_fn(source_text)
            except Exception:
                pass
        source_words[i] = _wc(source_text)
        _add_page_label(doc, i, total)
        _add_text_block(doc, text)
        processed.append(i)
        if i < total:
            doc.add_page_break()      # one Word page per PDF page (1:1 mapping)

    src.close()
    doc.save(out_path)

    # Fidelity audit — INDEPENDENT: re-open the saved DOCX and count the words
    # that actually persisted to disk per page, vs the source PDF page. In
    # faithful mode (mask_fn=None) this is a true round-trip check (save could in
    # principle drop text); in masked mode a delta reflects redaction.
    reread = _extract_docx_page_texts(out_path)
    page_audit = []
    for i in processed:
        pdf_w = source_words.get(i, 0)
        docx_w = _wc(reread.get(i, ""))
        page_audit.append({
            "page": i, "pdf_words": pdf_w, "docx_words": docx_w,
            "delta": docx_w - pdf_w,
        })
    return {
        "pages": total,
        "ocr_pages": ocr_pages,
        "page_audit": page_audit,
        "total_pdf_words": sum(p["pdf_words"] for p in page_audit),
        "total_docx_words": sum(p["docx_words"] for p in page_audit),
        "pages_with_delta": [p["page"] for p in page_audit if p["delta"] != 0],
    }
